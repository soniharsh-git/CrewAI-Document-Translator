from docx import Document

def breakdown_agent(input_doc_path):
    """
    Breakdown Agent:
    Reads the document and extracts paragraphs and tables into a structured format.

    Args:
        input_doc_path (str): Path to the input .docx file.

    Returns:
        list: Structured content list with paragraphs and tables.
    """
    document = Document(input_doc_path)
    structured_content = []

    for element in document.element.body:
        # Check if element is a paragraph
        if element.tag.endswith('p'):
            paragraph = element
            text = ''.join([node.text for node in paragraph.iter() if node.text])
            structured_content.append({
                "type": "paragraph",
                "text": text.strip(),
                "bold": False,      # We will improve this later
                "italic": False,    # We will improve this later
                "alignment": "left" # We can extract actual alignment later
            })

        # Check if element is a table
        elif element.tag.endswith('tbl'):
            table = Document().add_table(rows=0, cols=0)  # Temporary table object
            # Parse table content
            rows = []
            for row in element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'):
                cells = [cell.text for cell in row.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')]
                rows.append(cells)
            structured_content.append({
                "type": "table",
                "rows": rows
            })

    print("[Breakdown Agent] Document structured successfully.")
    return structured_content
