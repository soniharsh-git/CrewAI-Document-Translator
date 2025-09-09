from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def formatting_agent(translated_data, output_path="output/translated_document.docx"):
    """
    Formatting Agent:
    Reconstructs the translated structured data into a DOCX file.

    Args:
        translated_data (list): Translated structured data from translator agent.
        output_path (str): Path to save the translated DOCX file.

    Returns:
        None
    """
    doc = Document()

    for element in translated_data:
        if element['type'] == 'paragraph':
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element['text'])
            
            # Simulated styles (can be improved later)
            if element['bold']:
                run.bold = True
            if element['italic']:
                run.italic = True

            # Alignment (simplified)
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT
            }
            paragraph.alignment = alignment_map.get(element['alignment'], WD_PARAGRAPH_ALIGNMENT.LEFT)

        elif element['type'] == 'table':
            # Determine table size
            rows = element['rows']
            if not rows:
                continue
            
            table = doc.add_table(rows=0, cols=len(rows[0]))
            for row_data in rows:
                row = table.add_row().cells
                for idx, cell_text in enumerate(row_data):
                    row[idx].text = cell_text

    # Save the translated document
    doc.save(output_path)
    print(f"[Formatting Agent] Translated document saved to {output_path}")
